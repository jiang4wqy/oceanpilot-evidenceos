"""Build the OceanPilot 40-strong proposal DOCX from its Markdown source."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "submission" / "2026-08-16-oceanpilot-40-final-proposal.md"
OUTPUT = ROOT / "docs" / "submission" / "OceanPilot-40强完整参赛方案.docx"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x55, 0x55, 0x55)
LIGHT_FILL = "F4F6F9"
TABLE_FILL = "F4F6F9"
TABLE_BORDER = "D9E2EC"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, *, size: float | None = None, bold: bool | None = None,
                 italic: bool | None = None, color: RGBColor | None = None,
                 ascii_font: str = "Calibri", east_asia_font: str = "Microsoft YaHei") -> None:
    run.font.name = ascii_font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), ascii_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), ascii_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia_font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_inline(paragraph, text: str) -> None:
    """Render a compact subset of Markdown inline syntax."""
    token_re = re.compile(r"(\*\*.+?\*\*|`.+?`|<https?://[^>]+>|!\[[^\]]*\]\([^)]*\))")
    cursor = 0
    for match in token_re.finditer(text):
        if match.start() > cursor:
            set_run_font(paragraph.add_run(text[cursor:match.start()]))
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, ascii_font="Consolas", east_asia_font="Microsoft YaHei",
                         color=DARK_BLUE)
        elif token.startswith("<http"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, color=BLUE)
            run.underline = True
        cursor = match.end()
    if cursor < len(text):
        set_run_font(paragraph.add_run(text[cursor:]))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120,
                     bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_row_pagination(row, *, repeat_header: bool = False) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))
    if repeat_header and tr_pr.find(qn("w:tblHeader")) is None:
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        tr_pr.append(header)


def new_numbering_instance(doc: Document, style_name: str = "List Number") -> int:
    style_num_pr = doc.styles[style_name]._element.pPr.numPr
    base_num_id = int(style_num_pr.numId.val)
    numbering = doc.part.numbering_part.element
    base_num = next(
        node for node in numbering.findall(qn("w:num"))
        if int(node.get(qn("w:numId"))) == base_num_id
    )
    abstract_num_id = base_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    new_num_id = max(int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))) + 1
    new_num = OxmlElement("w:num")
    new_num.set(qn("w:numId"), str(new_num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_num_id)
    new_num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    new_num.append(override)
    numbering.append(new_num)
    return new_num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), "0")
    num_id_node = num_pr.find(qn("w:numId"))
    if num_id_node is None:
        num_id_node = OxmlElement("w:numId")
        num_pr.append(num_id_node)
    num_id_node.set(qn("w:val"), str(num_id))


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:color"), TABLE_BORDER)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("OceanPilot｜40 强完整参赛方案  ·  ")
    set_run_font(run, size=9, color=MUTED)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend((fld_char_begin, instr_text, fld_char_end))


def add_opening(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("鹿出草莓酱 × 钱海网络")
    set_run_font(r, size=12, bold=True, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(title)
    set_run_font(r, size=20, bold=True, color=RGBColor(0, 0, 0))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("40 强完整参赛方案｜2026-08-16｜Synthetic / Mock 数据边界")
    set_run_font(r, size=10.5, color=MUTED)


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_inline(p, text)


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    widths = [CONTENT_WIDTH_DXA // cols] * cols
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    if cols == 2:
        widths = [2700, 6660]
    elif cols == 3:
        widths = [1900, 3730, 3730]
    set_table_geometry(table, widths)
    set_table_borders(table)
    for row_idx, row in enumerate(rows):
        for col_idx in range(cols):
            cell = table.cell(row_idx, col_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_inline(p, row[col_idx] if col_idx < len(row) else "")
            if row_idx == 0:
                set_cell_shading(cell, TABLE_FILL)
                for run in p.runs:
                    run.bold = True
            elif cols > 2 and col_idx == 0:
                for run in p.runs:
                    run.bold = True
        set_row_pagination(table.rows[row_idx], repeat_header=row_idx == 0)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_image(doc: Document, alt: str, rel_path: str) -> None:
    image_path = SOURCE.parent / rel_path
    if not image_path.is_file():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[待补图片：{alt}]")
        set_run_font(run, italic=True, color=MUTED)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    picture = run.add_picture(str(image_path), width=Inches(6.35))
    picture._inline.docPr.set("descr", alt)
    picture._inline.docPr.set("title", alt)


def build() -> Path:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    title = lines[0].removeprefix("# ").strip()
    doc = Document()
    configure_styles(doc)
    add_page_field(doc.sections[0].footer.paragraphs[0])
    add_opening(doc, title)

    i = 1
    in_code = False
    code_lines: list[str] = []
    ordered_num_id: int | None = None
    while i < len(lines):
        line = lines[i]
        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                table = doc.add_table(rows=1, cols=1)
                set_table_geometry(table, [CONTENT_WIDTH_DXA])
                set_table_borders(table)
                cell = table.cell(0, 0)
                set_cell_shading(cell, "F7F9FC")
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                run = p.add_run("\n".join(code_lines))
                set_run_font(run, size=9.5, ascii_font="Consolas",
                             east_asia_font="Microsoft YaHei", color=DARK_BLUE)
                in_code = False
            ordered_num_id = None
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        if not line.strip():
            ordered_num_id = None
            i += 1
            continue

        image_match = re.fullmatch(r"!\[([^\]]*)\]\(([^)]+)\)", line.strip())
        if image_match:
            ordered_num_id = None
            add_image(doc, image_match.group(1), image_match.group(2))
            i += 1
            continue

        if line.startswith("|"):
            ordered_num_id = None
            raw_rows: list[list[str]] = []
            while i < len(lines) and lines[i].startswith("|"):
                cells = [cell.strip() for cell in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
                    raw_rows.append(cells)
                i += 1
            add_markdown_table(doc, raw_rows)
            continue

        heading = re.match(r"^(#{2,4})\s+(.+)$", line)
        if heading:
            ordered_num_id = None
            level = min(len(heading.group(1)) - 1, 3)
            p = doc.add_paragraph(style=f"Heading {level}")
            add_inline(p, heading.group(2))
            i += 1
            continue

        if line.startswith("> "):
            ordered_num_id = None
            block_lines = []
            while i < len(lines) and lines[i].startswith("> "):
                block_lines.append(lines[i][2:].strip())
                i += 1
            add_callout(doc, " ".join(block_lines))
            continue

        unordered = re.match(r"^-\s+(.+)$", line)
        ordered = re.match(r"^\d+\.\s+(.+)$", line)
        if unordered or ordered:
            p = doc.add_paragraph(style="List Bullet" if unordered else "List Number")
            if ordered:
                if ordered_num_id is None:
                    ordered_num_id = new_numbering_instance(doc)
                apply_numbering(p, ordered_num_id)
            else:
                ordered_num_id = None
            add_inline(p, (unordered or ordered).group(1))
            i += 1
            continue

        ordered_num_id = None
        p = doc.add_paragraph()
        add_inline(p, line)
        i += 1

    doc.core_properties.title = "OceanPilot：证据驱动的跨境商户成功智能体"
    doc.core_properties.subject = "AI 先锋未来人才大赛 40 强完整参赛方案"
    doc.core_properties.author = "鹿出草莓酱"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
